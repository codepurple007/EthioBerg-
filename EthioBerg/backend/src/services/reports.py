"""Composes pre-review analysis reports and renders them to DOCX.

A report is a rendering of an evaluation that already exists, not a second
opinion: it reuses the rule engine's results verbatim and adds the provenance a
reader needs to check them — which rule version was applied, which clause each
threshold came from, which page each figure was read off, and what the result
cannot be relied upon for.
"""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from src.domain.enums import RequirementState
from src.domain.models import (
    ActorRef,
    DocumentEvaluateResponse,
    IssuerDocument,
    ReportCandidate,
    ReportCaveat,
    ReportCitation,
    ReportEvidence,
    ReportPreview,
    RuleDefinition,
)
from src.services.rule_engine import PRE_REVIEW_DISCLAIMER, format_number

# Below this, a figure was matched loosely enough that a human should re-read it
# in the source document before relying on the requirement it fed.
LOW_CONFIDENCE = 0.8


def _now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _state_caveats(summary: dict[str, int]) -> list[ReportCaveat]:
    """Turn unresolved requirement states into plainly worded limitations."""
    caveats: list[ReportCaveat] = []

    conflicts = summary.get(RequirementState.CONFLICT, 0)
    if conflicts:
        caveats.append(
            ReportCaveat(
                severity="critical",
                message=(
                    f"{conflicts} requirement(s) had conflicting figures in the document. "
                    "These were not assessed, and the conflict must be resolved with the "
                    "issuer before any conclusion is drawn."
                ),
            )
        )

    missing = summary.get(RequirementState.MISSING_EVIDENCE, 0)
    if missing:
        caveats.append(
            ReportCaveat(
                severity="warning",
                message=(
                    f"{missing} requirement(s) could not be assessed because the figure they "
                    "depend on was not found in the document. Absence of evidence is not "
                    "evidence that the requirement is unmet."
                ),
            )
        )

    review = summary.get(RequirementState.PROFESSIONAL_REVIEW, 0)
    if review:
        caveats.append(
            ReportCaveat(
                severity="warning",
                message=(
                    f"{review} requirement(s) rest on draft rules that an administrator has not "
                    "yet approved, so they were skipped rather than evaluated."
                ),
            )
        )

    not_applicable = summary.get(RequirementState.NOT_APPLICABLE, 0)
    if not_applicable:
        caveats.append(
            ReportCaveat(
                severity="info",
                message=(
                    f"{not_applicable} rule(s) were not in force on the evaluation date and were "
                    "excluded."
                ),
            )
        )

    return caveats


def _evidence_caveats(evidence: list[ReportEvidence]) -> list[ReportCaveat]:
    """Flag figures whose extraction a reader should not take on trust."""
    caveats: list[ReportCaveat] = []

    low = [item.field for item in evidence if item.confidence < LOW_CONFIDENCE]
    if low:
        caveats.append(
            ReportCaveat(
                severity="warning",
                message=(
                    f"{len(low)} figure(s) were extracted with low confidence "
                    f"({', '.join(sorted(set(low)))}). Verify each against the source document."
                ),
            )
        )

    untraceable = [item.field for item in evidence if item.source_page is None]
    if untraceable:
        caveats.append(
            ReportCaveat(
                severity="warning",
                message=(
                    f"{len(untraceable)} figure(s) could not be traced to a page "
                    f"({', '.join(sorted(set(untraceable)))}), so the quotations below do not "
                    "cover them."
                ),
            )
        )

    return caveats


class ReportService:
    def __init__(self, repository, document_service):
        self.repository = repository
        self.document_service = document_service

    def list_candidates(self) -> list[ReportCandidate]:
        candidates: list[ReportCandidate] = []
        for document in self.document_service.list_documents():
            blocked = self._blocked_reason(document)
            candidates.append(
                ReportCandidate(
                    documentId=document.id,
                    filename=document.filename,
                    segment=document.segment,
                    uploadTimestamp=document.upload_timestamp,
                    extractionStatus=document.extraction_status,
                    factsConfirmed=document.facts_confirmed,
                    factCount=len(document.facts),
                    ready=not blocked,
                    blockedReason=blocked,
                )
            )
        return candidates

    @staticmethod
    def _blocked_reason(document: IssuerDocument) -> str:
        if not document.facts:
            return "No figures have been extracted from this document yet."
        if not document.facts_confirmed:
            return "Extracted figures must be confirmed under Document Review first."
        return ""

    def build_preview(self, document_id: str, actor: ActorRef) -> ReportPreview:
        document = self.document_service.get_document(document_id)
        evaluation = self.document_service.evaluation_for(document_id)

        evidence = self._evidence(document, evaluation)
        preview = ReportPreview(
            documentId=document.id,
            filename=document.filename,
            segment=document.segment,
            checksum=document.checksum,
            pageCount=document.page_count,
            uploadTimestamp=document.upload_timestamp,
            ruleVersion=evaluation.rule_version,
            generatedAt=_now_iso(),
            generatedBy=actor.actor_name,
            summary=evaluation.summary,
            categorySummary=evaluation.category_summary,
            requirements=evaluation.results,
            evidence=evidence,
            citations=self._citations(document, evaluation),
            caveats=[*_state_caveats(evaluation.summary), *_evidence_caveats(evidence)],
            disclaimer=PRE_REVIEW_DISCLAIMER,
        )
        self.repository.log_audit(actor, "REPORT_PREVIEWED", "IssuerDocument", document_id)
        return preview

    def _evidence(
        self, document: IssuerDocument, evaluation: DocumentEvaluateResponse
    ) -> list[ReportEvidence]:
        """Only the facts the evaluated rules actually consumed."""
        rules_by_id = {rule.rule_id: rule for rule in self._segment_rules(document)}
        used_fields = {
            rules_by_id[result.rule_id].field
            for result in evaluation.results
            if result.rule_id in rules_by_id
        }
        return [
            ReportEvidence(
                field=fact.field,
                value=fact.value,
                unit=fact.unit,
                period=fact.period,
                sourcePage=fact.source_page,
                sourceQuote=fact.source_quote,
                confidence=fact.confidence,
                status=fact.status,
            )
            for fact in document.facts
            if fact.field in used_fields
        ]

    def _citations(
        self, document: IssuerDocument, evaluation: DocumentEvaluateResponse
    ) -> list[ReportCitation]:
        rules_by_id = {rule.rule_id: rule for rule in self._segment_rules(document)}
        sources = {source.id: source for source in self.repository.get_sources()}

        citations: list[ReportCitation] = []
        for result in evaluation.results:
            rule = rules_by_id.get(result.rule_id)
            if rule is None:
                continue
            source = sources.get(rule.source_document_id)
            citations.append(
                ReportCitation(
                    ruleId=rule.rule_id,
                    ruleName=rule.name,
                    section=rule.source_section,
                    sourceTitle=source.title if source else rule.source_document_id,
                    issuingBody=source.issuing_body if source else "Unknown issuing body",
                    sourceVersion=source.version if source else "unknown",
                    publicationDate=source.publication_date if source else "",
                    url=source.url if source else "",
                )
            )
        return citations

    def _segment_rules(self, document: IssuerDocument) -> list[RuleDefinition]:
        return self.repository.rule_engine.list_rules(document.segment)

    def render_docx(self, document_id: str, actor: ActorRef) -> tuple[str, bytes]:
        """Return a suggested filename and the DOCX bytes."""
        preview = self.build_preview(document_id, actor)
        buffer = BytesIO()
        _build_document(preview).save(buffer)
        self.repository.log_audit(actor, "REPORT_EXPORTED", "IssuerDocument", document_id)
        stem = preview.filename.rsplit(".", 1)[0] or preview.document_id
        return f"{stem}-pre-review-report.docx", buffer.getvalue()


def _build_document(preview: ReportPreview):
    """Render a preview into a DOCX in the order a reviewer reads it."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Pre-Review Listing Readiness Analysis", level=0)

    subtitle = doc.add_paragraph(preview.filename)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.bold = True

    # The disclaimer leads, because a reader who stops after the first page must
    # still have seen it.
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(preview.disclaimer)
    disclaimer_run.bold = True
    disclaimer_run.font.size = Pt(10)

    doc.add_heading("Report details", level=1)
    _key_value_table(
        doc,
        [
            ("Document", preview.filename),
            ("Document ID", preview.document_id),
            ("Checksum (SHA-256)", preview.checksum),
            ("Pages", str(preview.page_count)),
            ("Market segment", str(preview.segment)),
            ("Rule version applied", preview.rule_version),
            ("Uploaded", preview.upload_timestamp),
            ("Generated", preview.generated_at),
            ("Generated by", preview.generated_by),
        ],
    )

    doc.add_heading("Summary", level=1)
    if preview.summary:
        _key_value_table(
            doc,
            [(_state_label(state), str(count)) for state, count in sorted(preview.summary.items())],
            headers=("Outcome", "Requirements"),
        )
    else:
        doc.add_paragraph("No requirements were evaluated.")

    doc.add_heading("Requirements", level=1)
    if preview.requirements:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        _header_row(table, ("Requirement", "Outcome", "Value found", "Threshold", "Basis"))
        for result in preview.requirements:
            cells = table.add_row().cells
            cells[0].text = result.rule_name
            cells[1].text = _state_label(result.state)
            cells[2].text = _format_value(result.fact_value)
            cells[3].text = result.threshold
            cells[4].text = result.source_section
    else:
        doc.add_paragraph("No requirements were evaluated.")

    doc.add_heading("Evidence", level=1)
    if preview.evidence:
        doc.add_paragraph(
            "Each figure below is quoted from the uploaded document at the page shown."
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        _header_row(table, ("Figure", "Value", "Period", "Page", "Quotation"))
        for item in preview.evidence:
            cells = table.add_row().cells
            cells[0].text = item.field
            cells[1].text = f"{_format_value(item.value)} {item.unit}".strip()
            cells[2].text = item.period or "—"
            cells[3].text = str(item.source_page) if item.source_page is not None else "not traced"
            cells[4].text = item.source_quote or "—"
    else:
        doc.add_paragraph("No figures were traced to the source document.")

    doc.add_heading("Citations", level=1)
    if preview.citations:
        for citation in preview.citations:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{citation.rule_name} — ").bold = True
            detail = f"{citation.source_title} ({citation.issuing_body}), {citation.section}"
            if citation.source_version:
                detail += f", version {citation.source_version}"
            if citation.publication_date:
                detail += f", published {citation.publication_date}"
            paragraph.add_run(detail)
            if citation.url:
                paragraph.add_run(f" — {citation.url}")
    else:
        doc.add_paragraph("No citations are available for the rules applied.")

    doc.add_heading("Caveats and limitations", level=1)
    if preview.caveats:
        for caveat in preview.caveats:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(f"[{caveat.severity.upper()}] ").bold = True
            paragraph.add_run(caveat.message)
    else:
        doc.add_paragraph(
            "Every requirement was assessed against a confirmed figure traced to the document."
        )

    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(preview.disclaimer)
    doc.add_paragraph(
        "This report was produced automatically from the uploaded document and the rule "
        f"version named above ({preview.rule_version}). It records how those rules compare "
        "with figures read out of that document. It is not advice, and it does not constitute "
        "a submission to, or a decision by, any regulator or exchange."
    )

    return doc


def _state_label(state: str) -> str:
    return str(state).replace("_", " ").capitalize()


def _format_value(value: float | int | str | None) -> str:
    if value is None:
        return "—"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (int, float)):
        return format_number(value)
    return str(value)


def _header_row(table, headers: tuple[str, ...]) -> None:
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _key_value_table(doc, rows: list[tuple[str, str]], headers: tuple[str, str] | None = None):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    _header_row(table, headers or ("Field", "Value"))
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    return table
