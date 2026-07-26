from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from src.services.ocr import RENDER_DPI, OcrOptions, image_to_text, probe, resolve_languages

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024


@dataclass
class ParsedPage:
    page_number: int
    text: str
    ocr_applied: bool = False


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]
    mime_type: str
    checksum: str
    ocr_pages: list[int] = field(default_factory=list)
    ocr_error: str = ""

    @property
    def full_text(self) -> str:
        return "\n".join(page.text for page in self.pages)


class IngestionError(Exception):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code
        self.message = message


def compute_checksum(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def validate_upload(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise IngestionError(
            "UNSUPPORTED_TYPE",
            "Supported issuer documents: PDF (primary) and DOCX.",
        )
    if len(content) == 0:
        raise IngestionError("EMPTY_FILE", "Uploaded file is empty.")
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise IngestionError("FILE_TOO_LARGE", "File exceeds the 25 MB upload limit.")
    if suffix == ".pdf" and not content.startswith(b"%PDF"):
        raise IngestionError("INVALID_PDF", "File extension is PDF but content is not a valid PDF.")
    if suffix == ".docx" and content[:2] != b"PK":
        raise IngestionError("INVALID_DOCX", "File extension is DOCX but content is not a valid DOCX archive.")
    return suffix


def parse_pdf(content: bytes, ocr: OcrOptions | None = None) -> tuple[list[ParsedPage], list[int], str]:
    """Extract text per page, falling back to OCR on pages that carry no text layer.

    Returns the pages, the page numbers OCR recovered, and a reason string when OCR
    was wanted but could not run.
    """
    import fitz

    options = ocr or OcrOptions()
    pages: list[ParsedPage] = []
    ocr_pages: list[int] = []
    ocr_error = ""
    languages: list[str] | None = None

    with fitz.open(stream=content, filetype="pdf") as doc:
        for index, page in enumerate(doc, start=1):
            text = page.get_text("text")
            needs_ocr = options.enabled and len(text.strip()) < options.min_text_chars

            if needs_ocr and not ocr_error:
                if languages is None:
                    status = probe()
                    if not status.available:
                        ocr_error = status.detail
                    else:
                        languages = resolve_languages(options.languages)
                        if not languages:
                            ocr_error = (
                                "None of the configured OCR languages "
                                f"({', '.join(options.languages)}) are installed."
                            )

            if needs_ocr and not ocr_error and languages:
                try:
                    image = page.get_pixmap(dpi=RENDER_DPI).tobytes("png")
                    recovered = image_to_text(image, languages).strip()
                except Exception as exc:  # noqa: BLE001 - one bad page must not fail the upload
                    ocr_error = f"OCR failed on page {index}: {exc}"
                else:
                    if recovered:
                        pages.append(ParsedPage(page_number=index, text=recovered, ocr_applied=True))
                        ocr_pages.append(index)
                        continue

            pages.append(ParsedPage(page_number=index, text=text))

    return pages, ocr_pages, ocr_error


def parse_docx(content: bytes) -> list[ParsedPage]:
    from docx import Document

    document = Document(BytesIO(content))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return [ParsedPage(page_number=1, text=text)]


def parse_document(
    filename: str, content: bytes, ocr: OcrOptions | None = None
) -> ParsedDocument:
    suffix = validate_upload(filename, content)
    checksum = compute_checksum(content)
    ocr_pages: list[int] = []
    ocr_error = ""
    if suffix == ".pdf":
        pages, ocr_pages, ocr_error = parse_pdf(content, ocr)
        mime_type = "application/pdf"
    else:
        # DOCX always carries a text layer, so OCR never applies.
        pages = parse_docx(content)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return ParsedDocument(
        pages=pages,
        mime_type=mime_type,
        checksum=checksum,
        ocr_pages=ocr_pages,
        ocr_error=ocr_error,
    )


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()
