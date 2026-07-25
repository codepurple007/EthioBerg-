from docx import Document

output_path = "sample_issuer_prospectus.docx"
doc = Document()
doc.add_heading("Draft Prospectus — Demo Issuer S.C.", level=1)
doc.add_paragraph(
    "This sample issuer document is provided for EthioBerg listing readiness demonstration only."
)
doc.add_paragraph(
    "The issuer has an operating track record of 4 years in its principal business activities."
)
doc.add_paragraph(
    "Based on the latest valuation, estimated market capitalization ETB 620 million is presented in this section."
)
doc.add_paragraph(
    "The public free float represents 14 percent of the total issued ordinary shares."
)
doc.add_paragraph("The issuer has 145 shareholders at the reporting date.")
doc.add_paragraph("Financial statements are prepared in accordance with IFRS and received an unqualified audit opinion.")
doc.save(output_path)
print(f"Wrote {output_path}")
