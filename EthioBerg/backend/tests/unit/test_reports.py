"""Report composition and DOCX rendering.

The report is a legal-adjacent artefact, so these tests are mostly about what it
must never do: overstate an unassessed requirement, drop a caveat, or omit the
disclaimer.
"""

from __future__ import annotations

from io import BytesIO

import pytest

from src.domain.enums import RequirementState
from src.domain.models import ReportCaveat, ReportEvidence
from src.services.reports import (
    LOW_CONFIDENCE,
    _build_document,
    _evidence_caveats,
    _format_value,
    _state_caveats,
)
from src.services.rule_engine import PRE_REVIEW_DISCLAIMER, format_number


def _severities(caveats: list[ReportCaveat]) -> set[str]:
    return {caveat.severity for caveat in caveats}


def test_conflicting_requirements_raise_a_critical_caveat():
    caveats = _state_caveats({RequirementState.CONFLICT: 2})

    assert "critical" in _severities(caveats)
    assert "2 requirement(s)" in caveats[0].message


def test_missing_evidence_is_not_reported_as_a_failure():
    caveats = _state_caveats({RequirementState.MISSING_EVIDENCE: 3})

    message = caveats[0].message
    assert "could not be assessed" in message
    assert "Absence of evidence is not evidence that the requirement is unmet" in message


def test_draft_rules_are_declared_as_skipped():
    caveats = _state_caveats({RequirementState.PROFESSIONAL_REVIEW: 1})

    assert "draft rules" in caveats[0].message
    assert "skipped rather than evaluated" in caveats[0].message


def test_a_clean_evaluation_produces_no_caveats():
    assert _state_caveats({RequirementState.MET: 5}) == []


def _evidence(**overrides) -> ReportEvidence:
    defaults = {
        "field": "market_cap_etb",
        "value": 620_000_000,
        "unit": "ETB",
        "period": None,
        "sourcePage": 18,
        "sourceQuote": "Estimated market capitalization ETB 620 million.",
        "confidence": 0.95,
        "status": "USER_CONFIRMED",
    }
    return ReportEvidence(**{**defaults, **overrides})


def test_low_confidence_figures_are_flagged():
    caveats = _evidence_caveats([_evidence(confidence=LOW_CONFIDENCE - 0.01)])

    assert "low confidence" in caveats[0].message
    assert "market_cap_etb" in caveats[0].message


def test_confident_traced_figures_are_not_flagged():
    assert _evidence_caveats([_evidence()]) == []


def test_untraced_figures_are_flagged():
    caveats = _evidence_caveats([_evidence(sourcePage=None)])

    assert any("could not be traced to a page" in caveat.message for caveat in caveats)


@pytest.mark.parametrize(
    ("value", "expected"),
    [
        (500_000_000, "500,000,000"),
        (3, "3"),
        (14.5, "14.50"),
        (None, "—"),
        ("n/a", "n/a"),
    ],
)
def test_values_are_rendered_for_a_reader_not_a_debugger(value, expected):
    assert _format_value(value) == expected


def test_large_thresholds_avoid_scientific_notation():
    assert format_number(500_000_000) == "500,000,000"
    assert "e+" not in format_number(620_000_000)


def _preview(**overrides):
    from src.domain.models import ReportPreview, RequirementResult

    defaults = dict(
        documentId="doc-1",
        filename="prospectus.pdf",
        segment="MAIN",
        checksum="abc123",
        pageCount=21,
        uploadTimestamp="2026-07-20T09:15:00+00:00",
        ruleVersion="2025.1-draft",
        generatedAt="2026-07-26T01:00:00+00:00",
        generatedBy="Test Admin",
        summary={RequirementState.MET: 1, RequirementState.CONFLICT: 1},
        categorySummary=[],
        requirements=[
            RequirementResult(
                ruleId="ESX_MAIN_MARKET_CAP",
                ruleName="Minimum market capitalization",
                state=RequirementState.MET,
                factValue=620_000_000,
                threshold="≥ 500,000,000 ETB",
                category="Size",
                sourceSection="Volume C, Main Market listing criteria",
            )
        ],
        evidence=[_evidence()],
        citations=[],
        caveats=[ReportCaveat(severity="critical", message="Conflicting figures were found.")],
        disclaimer=PRE_REVIEW_DISCLAIMER,
    )
    return ReportPreview(**{**defaults, **overrides})


def _docx_text(preview) -> str:
    from docx import Document

    buffer = BytesIO()
    _build_document(preview).save(buffer)
    buffer.seek(0)
    document = Document(buffer)

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_docx_always_carries_the_disclaimer():
    assert PRE_REVIEW_DISCLAIMER in _docx_text(_preview())


def test_docx_names_the_rule_version_it_applied():
    assert "2025.1-draft" in _docx_text(_preview())


def test_docx_carries_every_caveat():
    text = _docx_text(_preview())

    assert "Conflicting figures were found." in text
    assert "CRITICAL" in text


def test_docx_quotes_evidence_with_its_page():
    text = _docx_text(_preview())

    assert "Estimated market capitalization ETB 620 million." in text
    assert "18" in text


def test_docx_states_when_a_figure_was_not_traced():
    text = _docx_text(_preview(evidence=[_evidence(sourcePage=None)]))

    assert "not traced" in text


def test_docx_renders_large_figures_readably():
    text = _docx_text(_preview())

    assert "620,000,000" in text
    assert "6.2e+08" not in text
